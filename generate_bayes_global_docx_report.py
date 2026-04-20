#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Gera relatório em Word (.docx) equivalente ao PDF do BayesGlobal.

Saída:
  analysis_proba_raw/pro_portfolio_all/Relatorio_BayesGlobal_Mesa_Profissional_<YYYY-MM-DD>.docx

Fontes:
  - oos_walkforward_global_bayes_weekly.csv
  - oos_walkforward_global_bayes_daily.csv
  - oos_walkforward_global_bayes_selected_rules.csv
  - portfolio_refined_global_bayes_full_comparison.csv
  - scored_dedup_proba_raw_all.csv (para métricas por combinação)
"""

from __future__ import annotations

import math
from datetime import date
from pathlib import Path
from typing import Dict, List

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


OUT_DIR = Path("/workspace/analysis_proba_raw/pro_portfolio_all")
SCORED = Path("/workspace/analysis_proba_raw/scored_dedup_proba_raw_all.csv")

WF_WEEKLY = OUT_DIR / "oos_walkforward_global_bayes_weekly.csv"
WF_DAILY = OUT_DIR / "oos_walkforward_global_bayes_daily.csv"
WF_RULES = OUT_DIR / "oos_walkforward_global_bayes_selected_rules.csv"
COMPARISON = OUT_DIR / "portfolio_refined_global_bayes_full_comparison.csv"

# parâmetros exibidos (devem refletir a versão atual)
BANKROLL = 2300.0
MAX_DAILY_EXPOSURE_FRAC_Q = 0.70
MAX_DAILY_DRAWDOWN_FRAC = 0.25
MAX_P_DAILY_DD = 0.10

MIN_SELECTED_BETS = 6
MIN_NONZERO_WEEKS = 6
MIN_BETS_PER_BIN = 20
MIN_BINS_FOR_STABILITY = 3

HYSTERESIS_ENABLED = False
HYST_P_SWITCH = 0.90

ROBUST_CUTOFF_ENABLED = False
ROBUST_CUTOFF_DELTA = 0.02


def compute_weekly_stats(w: np.ndarray) -> Dict[str, float]:
    x = np.asarray(w, dtype=float)
    x = x[np.isfinite(x)]
    if x.size == 0:
        return {"n": 0}
    m = float(x.mean())
    s = float(x.std(ddof=1)) if x.size > 1 else 0.0
    med = float(np.median(x))
    pneg = float((x < 0).mean())
    sharpe_w = float(m / s) if s > 0 else float("nan")
    sharpe_ann = float(sharpe_w * math.sqrt(52.0)) if np.isfinite(sharpe_w) else float("nan")
    return {"n": int(x.size), "mean": m, "std": s, "median": med, "pneg": pneg, "sharpe_week": sharpe_w, "sharpe_annual": sharpe_ann}


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = None
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_h(doc: Document, level: int, text: str) -> None:
    doc.add_heading(text, level=level)


def add_kv_table(doc: Document, rows: List[List[str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = rows[0][0]
    hdr[1].text = rows[0][1]
    for r in rows[1:]:
        row = table.add_row().cells
        row[0].text = str(r[0])
        row[1].text = str(r[1])


def main() -> int:
    today = date.today().isoformat()
    out_path = OUT_DIR / f"Relatorio_BayesGlobal_Mesa_Profissional_{today}.docx"

    wf_week = pd.read_csv(WF_WEEKLY)
    wf_daily = pd.read_csv(WF_DAILY)
    wf_rules = pd.read_csv(WF_RULES)
    comp = pd.read_csv(COMPARISON) if COMPARISON.exists() else pd.DataFrame()

    # derive current-week rules from WF_RULES (consistent with audit)
    wf_rules["active"] = (wf_rules["status"] == "ok") & (wf_rules["stake_frac"] > 0)
    last_week = str(wf_rules["test_week"].iloc[-1])
    cur = wf_rules[(wf_rules["test_week"] == last_week) & (wf_rules["active"])].copy()
    cur_csv_path = OUT_DIR / "global_bayes_current_week_rules.csv"
    cur.to_csv(cur_csv_path, index=False)

    # aggregate metrics
    w = wf_week["profit_cap2_usd"].to_numpy(dtype=float)
    wstats = compute_weekly_stats(w)
    w_traded = wf_week.loc[wf_week["stake_usd"] > 0, "profit_cap2_usd"].to_numpy(dtype=float)
    wstats_t = compute_weekly_stats(w_traded)

    stake_tot = float(wf_week["stake_usd"].sum())
    profit_tot = float(wf_week["profit_cap2_usd"].sum())
    roi_on_stake = float(profit_tot / stake_tot) if stake_tot > 0 else float("nan")
    mean_week = float(wstats.get("mean", float("nan")))

    # daily risk
    s_day = wf_daily["stake_usd"].to_numpy(dtype=float) if not wf_daily.empty else np.array([])
    p_day = wf_daily["profit_cap2_usd"].to_numpy(dtype=float) if not wf_daily.empty else np.array([])
    p80_exp = float(np.quantile(s_day, 0.80)) if s_day.size else float("nan")
    var10 = float(np.quantile(p_day, 0.10)) if p_day.size else float("nan")
    p_dd = float(np.mean(p_day <= (-MAX_DAILY_DRAWDOWN_FRAC * BANKROLL))) if p_day.size else float("nan")

    # doc
    doc = Document()
    add_title(doc, "Relatório — Portfólio Bayes Global (mesa profissional)")
    doc.add_paragraph(f"Data: {today}")
    doc.add_paragraph("Workspace: /workspace")

    add_h(doc, 2, "1. O que foi feito (método e objetivo)")
    doc.add_paragraph(
        "Construímos e validamos uma estratégia de apostas por portfólio de combinações "
        "(DoW × tipo FT/FH × score ≥ cutoff × stake). O score é proba_raw (alinhado ao processo operacional). "
        "A avaliação OOS é via walk-forward semanal, com constraints de risco diário e controle global por α."
    )

    add_h(doc, 2, "6. Parâmetros e constraints usados")
    add_kv_table(
        doc,
        [
            ["Item", "Valor"],
            ["Banca (BANKROLL)", f"USD {BANKROLL:,.0f}"],
            ["Exposição diária", f"p80(soma stakes/dia) ≤ {MAX_DAILY_EXPOSURE_FRAC_Q*100:.0f}% da banca"],
            ["Risco diário", f"VaR10% do PnL diário ≥ -{MAX_DAILY_DRAWDOWN_FRAC*100:.0f}% da banca; P(loss≥25%) ≤ {MAX_P_DAILY_DD*100:.0f}%"],
            ["Confiança mínima (apostas)", f"MIN_SELECTED_BETS={MIN_SELECTED_BETS}; MIN_NONZERO_WEEKS={MIN_NONZERO_WEEKS}"],
            ["Confiança mínima (bins)", f"MIN_BETS_PER_BIN={MIN_BETS_PER_BIN}; exigir ≥{MIN_BINS_FOR_STABILITY} bins"],
            ["Histerese", f"habilitado={HYSTERESIS_ENABLED}; P(switch)≥{HYST_P_SWITCH:.2f}"],
            ["Robust cutoff", f"habilitado={ROBUST_CUTOFF_ENABLED}; vizinhança ±{ROBUST_CUTOFF_DELTA:.2f}"],
        ],
    )

    add_h(doc, 2, "2. Portfólio otimizado escolhido (Bayes Global)")
    doc.add_paragraph(
        "O Bayes Global é uma política dinâmica (as regras mudam ao longo do tempo). "
        f"Abaixo estão as regras da semana mais recente disponível no dataset ({last_week})."
    )
    if cur.empty:
        doc.add_paragraph("Nenhuma combinação ativa na última semana do walk-forward.")
    else:
        cur2 = cur.sort_values(["bet_type", "dow_pt"])
        table = doc.add_table(rows=1, cols=6)
        hdr = table.rows[0].cells
        hdr[0].text = "Tipo"
        hdr[1].text = "Dia"
        hdr[2].text = "Score"
        hdr[3].text = "Cutoff"
        hdr[4].text = "Stake"
        hdr[5].text = "α"
        for _, r in cur2.iterrows():
            row = table.add_row().cells
            row[0].text = str(r["bet_type"])
            row[1].text = str(r["dow_pt"])
            row[2].text = str(r["score_col"])
            row[3].text = f"{float(r['cutoff']):.2f}"
            row[4].text = f"{float(r['stake_frac'])*100:.1f}%"
            row[5].text = f"{float(r['alpha_global']):.3f}"

    add_h(doc, 2, "3. Métricas estatísticas e de negócio")
    add_h(doc, 3, "3.1 Bayes Global (walk-forward OOS, cap2)")
    add_kv_table(
        doc,
        [
            ["Métrica", "Valor"],
            ["Semanas (WF)", str(int(wstats["n"]))],
            ["Lucro total (WF)", f"USD {profit_tot:,.1f}"],
            ["Stake total (WF)", f"USD {stake_tot:,.1f}"],
            ["ROI por $ apostado (WF)", f"{roi_on_stake:.4f}"],
            ["Lucro médio semanal", f"USD {mean_week:,.1f}"],
            ["Mediana semanal", f"USD {wstats['median']:,.1f}"],
            ["Std semanal", f"USD {wstats['std']:,.1f}"],
            ["P(semana<0)", f"{wstats['pneg']*100:.1f}%"],
            ["Sharpe anualizado", f"{wstats['sharpe_annual']:.3f}"],
            ["— Apenas semanas com trade (stake>0) —", ""],
            ["Lucro médio semanal (trade)", f"USD {wstats_t['mean']:,.1f}" if wstats_t["n"] else "nan"],
            ["Mediana semanal (trade)", f"USD {wstats_t['median']:,.1f}" if wstats_t["n"] else "nan"],
            ["Std semanal (trade)", f"USD {wstats_t['std']:,.1f}" if wstats_t["n"] else "nan"],
            ["P(semana<0 | trade)", f"{wstats_t['pneg']*100:.1f}%" if wstats_t["n"] else "nan"],
            ["Risco diário: p80 stake/dia", f"USD {p80_exp:,.0f} (limite USD {MAX_DAILY_EXPOSURE_FRAC_Q*BANKROLL:,.0f})"],
            ["Risco diário: VaR10% PnL/dia", f"USD {var10:,.1f} (limite ≥ USD {-MAX_DAILY_DRAWDOWN_FRAC*BANKROLL:,.0f})"],
            ["Risco diário: P(PnL dia ≤ -25%)", f"{p_dd*100:.1f}% (limite ≤ {MAX_P_DAILY_DD*100:.0f}%)"],
        ],
    )

    add_h(doc, 3, "3.2 Comparação com portfólio fixo in-sample")
    if comp.empty:
        doc.add_paragraph("Arquivo de comparação não encontrado.")
    else:
        cols = ["name", "weeks", "bets", "stake", "profit", "roi_on_stake", "mean_week", "std_week", "pneg_week"]
        t = doc.add_table(rows=1, cols=len(cols))
        for i, c in enumerate(cols):
            t.rows[0].cells[i].text = c
        for _, r in comp[cols].iterrows():
            row = t.add_row().cells
            for i, c in enumerate(cols):
                row[i].text = str(r[c])

    add_h(doc, 2, "Anexo A — Auditoria: segmentos ativos por semana (walk-forward)")
    doc.add_paragraph("Lista de segmentos ativos (status=ok e stake>0) por semana.")
    audit = []
    for wk, g in wf_rules.groupby("test_week", sort=False):
        ga = g[g["active"]].copy()
        segs = sorted(ga["rule_key"].astype(str).tolist())
        a = float(ga["alpha_global"].iloc[0]) if not ga.empty else float("nan")
        audit.append((str(wk), f"{a:.3f}" if np.isfinite(a) else "nan", str(len(segs)), ", ".join(segs)))
    t = doc.add_table(rows=1, cols=4)
    t.rows[0].cells[0].text = "Semana"
    t.rows[0].cells[1].text = "α"
    t.rows[0].cells[2].text = "N ativos"
    t.rows[0].cells[3].text = "Segmentos"
    for wk, a, n, segs in audit:
        row = t.add_row().cells
        row[0].text = wk
        row[1].text = a
        row[2].text = n
        row[3].text = segs

    doc.add_paragraph(f"Fontes: {WF_WEEKLY.name}, {WF_DAILY.name}, {WF_RULES.name}, global_bayes_current_week_rules.csv, {COMPARISON.name}")
    doc.save(str(out_path))
    print(str(out_path))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

